from openpyxl import Workbook
from docx import Document

"""
Dmitry Knyazhevskiy - COMP1112 Assignment 2
This program goes through each of the generated invoices one by one
Strips all the important data
Then saves the data to a spreadsheet
"""

wb = Workbook() #Create a new workbook called workbook
dest_filename = 'A2.xlsx' #Creates the file for that workbook
sheet1 = wb.active #Sets the first sheet as the active one
sheet1.title = "Sheet" #Titles the first sheet as Sheet

#Labels all the columns
sheet1['A1'] = "Invoice Number" 
sheet1['B1'] = "Total Quantity"
sheet1['C1'] = "Subtotal"
sheet1['D1'] = "Tax"
sheet1['E1'] = "Total"

for k in range(200): #200 Docs were created so it iterates from 0 - 200

    sample = Document("INV1000" + '%03d' % k + ".docx") #All of the docs start with INV1000 so we simply append the index of the iteration formatted to 3 digits

    sheet1.cell(row=k+2,column=1).value = "INV1000" + '%03d' % k # For the first row of every column it sets the value to the invoice number

    cleanSample = sample.paragraphs[1].text.replace(":"," ") #Grabs the first paragraph which contains the products and their amount, then replacing all : with spaces to separate the numbers from words
    nums = [int(i) for i in cleanSample.split() if i.isdigit()] #Goes through all the substrings and saves the numeric ones to a list called nums
    sheet1.cell(row=k+2,column=2).value = sum(nums) #Appends the sum of the numbers in the list to the second column

    cleanSample = sample.paragraphs[2].text.replace(":"," ") #Grabs the second paragraph which contains the products and their amount, then replacing all : with spaces to separate the numbers from words
    nums2 = [(i) for i in cleanSample.split()] #Since all paragraph twos follow the same format, splitting it makes 6 pieces in the list nums2 which correspond to the same values on every invoice
    sheet1.cell(row=k+2,column=3).value = nums2[1] # Index 1 of the nums2 list is the subtotal, which gets saved to third column of the current row
    sheet1.cell(row=k+2,column=4).value = nums2[3] # Index 3 of the nums2 list is the tax, which gets saved to fourth column of the current row
    sheet1.cell(row=k+2,column=5).value = nums2[5] # Index 5 of the nums2 list is the total, which gets saved to fifth column of the current row  

wb.save('A2.xlsx') # Saves the edited workbook to the A2 spreadsheet

